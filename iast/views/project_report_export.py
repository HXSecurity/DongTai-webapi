# coding:utf-8

import time
from collections import namedtuple
from iast.utils import get_model_field

from django.db.models import Q
from django.http import FileResponse
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from dongtai.models.vulnerablity import IastVulnerabilityModel
from dongtai.endpoint import R
from iast.base.agent import get_vul_count_by_agent
from dongtai.endpoint import UserEndPoint
from dongtai.models.agent import IastAgent
from dongtai.models.project import IastProject
from dongtai.models.vul_level import IastVulLevel
from webapi.settings import MEDIA_ROOT
from django.utils.translation import gettext_lazy as _
from rest_framework import serializers
from iast.utils import extend_schema_with_envcheck, get_response_serializer
from io import BytesIO
from dongtai.models.project_report import ProjectReport
from django.http import HttpResponse
import logging
import os
import xlwt
import xlrd

logger = logging.getLogger('dongtai-webapi')


class _ProjectReportExportQuerySerializer(serializers.Serializer):
    vid = serializers.CharField(
        help_text=_("The version id of the project"))
    pname = serializers.CharField(
        help_text=_("The name of the project"))
    pid = serializers.IntegerField(help_text=_("The id of the project"))


class ProjectReportExport(UserEndPoint):
    name = 'api-v1-word-maker'
    description = _('Vulnerability Report Generate - Word')

    @staticmethod
    def create_word():
        pass

    @staticmethod
    def get_agents_with_project_id(pid, auth_users):
        """
        :param pid:
        :param auth_users:
        :return:
        """
        relations = IastAgent.objects.filter(bind_project_id=pid,
                                             user__in=auth_users).values("id")
        agent_ids = [relation['id'] for relation in relations]
        return agent_ids

    @staticmethod
    def create_report():
        pass

    def get(self, request):
        # 生成报告
        timestamp = time.time()
        id = 0
        try:
            id = int(request.query_params.get("id", 0))
        except:
            pass

        report = ProjectReport.objects.filter(pk=id).first()

        if report:
            project_id = 0
            auth_users = self.get_auth_users(report.user)
            if report.project:
                project_id = report.project.id
            file_name, file_stream = self.generate_report(
                report.type,
                project_id,
                report.vul_id,
                auth_users,
                report.user,
                timestamp
            )

            # if file_name:
            #     report_filename = _('Vulnerability Report - {}. {}').format(
            #         timestamp, type)
            #     file_stream.seek(0)
            #     from wsgiref.util import FileWrapper
            #     response = FileResponse(FileWrapper(file_stream))
            #     response['content_type'] = 'application/octet-stream'
            #     response[
            #         'Content-Disposition'] = f"attachment; filename={report_filename}"
            #     return response

            if file_name:
                report_filename = _('Vulnerability Report - {}. {}').format(
                    timestamp, type)
                response = HttpResponse(file_stream)
                return response

        return R.failure(status=203, msg=_('no permission'))

    def generate_word_report(self, user, project, vul, count_result, levelInfo, timestamp):
        document = Document()
        document.styles.add_style('TitleOne', WD_STYLE_TYPE.PARAGRAPH).font.name = 'Arial'
        document.styles.add_style('TitleTwo', WD_STYLE_TYPE.PARAGRAPH).font.name = 'Arial'
        document.styles.add_style('TitleThree', WD_STYLE_TYPE.PARAGRAPH).font.name = 'Arial'
        document.styles.add_style('TitleFour', WD_STYLE_TYPE.PARAGRAPH).font.name = 'Arial'

        document.add_heading(u'%s' % project.name, 0)

        document.add_heading(u'%s' % project.mode, 2)

        timeArray = time.localtime(project.latest_time)
        otherStyleTime = time.strftime("%Y-%m-%d %H:%M:%S", timeArray)

        pTime = document.add_paragraph(u'%s' % otherStyleTime)
        pTime.paragraph_format.space_before = Pt(400)
        pTime.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        pReport = document.add_paragraph(_(u'Security Testing Report'))
        pReport.paragraph_format.line_spacing = Pt(20)
        pReport.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer = document.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.add_run(u'北京安全共识科技有限公司')
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_page_break()

        oneTitle = document.add_paragraph()
        oneTitle.add_run(_(u'First, project information')).font.name = 'Arial'
        oneTitle.style = "TitleOne"

        table = document.add_table(rows=1, cols=2, style='Table Grid')

        hdr_cells = table.rows[0].cells

        new_cells = table.add_row().cells
        new_cells[0].text = _('Application name')
        new_cells[1].text = project.name
        new_cells = table.add_row().cells
        new_cells[0].text = _('Author')
        new_cells[1].text = user.username
        new_cells = table.add_row().cells
        new_cells[0].text = _('Application type')
        new_cells[1].text = project.mode
        new_cells = table.add_row().cells
        new_cells[0].text = _('Number of Vulnerability')
        new_cells[1].text = str(project.vul_count)
        new_cells = table.add_row().cells
        new_cells[0].text = _('Number of Agent')
        new_cells[1].text = str(project.agent_count)
        new_cells = table.add_row().cells
        new_cells[0].text = _('Latest time')
        new_cells[1].text = time.strftime("%Y-%m-%d %H:%M:%S", timeArray)

        levelNameArr = {}
        levelIdArr = {}
        if levelInfo:
            for level_item in levelInfo:
                levelNameArr[level_item.name_value] = level_item.id
                levelIdArr[level_item.id] = level_item.name_value

        type_summary = count_result['type_summary']

        levelCount = count_result['levelCount']

        vulDetail = count_result['vulDetail']

        oneTitle = document.add_paragraph()
        oneTitle.add_run(_(u'Second, the result analysis'))
        oneTitle.style = "TitleOne"

        twoTitle = document.add_paragraph()
        twoTitle.add_run(_(u'2.1 Vulnerability Severity Levels Distribution'))
        twoTitle.style = "TitleTwo"
        levelCountArr = []
        if levelCount:
            for ind in levelCount.keys():
                levelCountArr.append(str(levelIdArr[ind]) + str(levelCount[ind]))
        levelCountStr = ",".join(levelCountArr)
        document.add_paragraph(levelCountStr)

        twoTitle = document.add_paragraph()
        twoTitle.add_run(_(u'2.2 Distribution of Vulnerability'))
        twoTitle.style = "TitleTwo"

        table = document.add_table(rows=1, cols=3, style='Table Grid')

        hdr_cells = table.rows[0].cells

        hdr_cells[0].text = _('Severity levels')
        hdr_cells[1].text = _('Vulnerability type name')
        hdr_cells[2].text = _('Number')
        if type_summary:
            for type_item in type_summary:
                new_cells = table.add_row().cells
                new_cells[0].text = levelIdArr[type_item['type_level']]
                new_cells[1].text = type_item['type_name']
                new_cells[2].text = str(type_item['type_count'])

        document.add_page_break()

        twoTitle = document.add_paragraph()
        twoTitle.add_run(_(u'2.3 Vulnerability details'))
        twoTitle.style = "TitleTwo"

        if vulDetail:
            type_ind = 1
            for vul in vulDetail.keys():

                threeTitle = document.add_paragraph()
                threeTitle.add_run(u'%s(%s)' % ("2.3." + str(type_ind) + "  " + vul, len(vulDetail[vul])))
                threeTitle.style = "TitleThree"
                if vulDetail[vul]:
                    ind = 1
                    for one in vulDetail[vul]:
                        p = document.add_paragraph()
                        p.add_run("2.3." + str(type_ind) + "." + str(ind) + "  " + one['title']).bold = True
                        p.style = "TitleFour"
                        ind = ind + 1
                        document.add_heading(_(u'Summary'), level=4)

                        table = document.add_table(rows=1, cols=2, style='Table Grid')
                        new_cells = table.add_row().cells
                        new_cells[0].text = _("Severity level")
                        new_cells[1].text = levelIdArr[one['level_id']]

                        new_cells = table.add_row().cells
                        new_cells[0].text = _("First scan time")
                        new_cells[1].text = one['first_time']

                        new_cells = table.add_row().cells
                        new_cells[0].text = _("Last scan time")
                        new_cells[1].text = one['latest_time']

                        new_cells = table.add_row().cells
                        new_cells[0].text = _("Development language")
                        new_cells[1].text = one['language']

                        new_cells = table.add_row().cells
                        new_cells[0].text = _("Vulnerability URL")
                        new_cells[1].text = one['url']
                        document.add_heading(_(u'Vulnerability description'), level=4)
                        if one['detail_data']:
                            for item in one['detail_data']:
                                document.add_paragraph(u'%s' % item)
                type_ind = type_ind + 1

        document.styles['TitleOne'].font.size = Pt(20)
        document.styles['TitleOne'].font.name = "Arial"
        document.styles['TitleTwo'].font.size = Pt(18)
        document.styles['TitleTwo'].font.name = "Arial"
        document.styles['TitleThree'].font.size = Pt(16)
        document.styles['TitleFour'].font.size = Pt(14)
        filename = f"{MEDIA_ROOT}/reports/vul-report-{user.id}-{timestamp}.docx"
        file_stream = BytesIO()
        document.save(file_stream)
        return filename, file_stream

    def generate_report(self, type, pid, vid, auth_users, user,
                        timestamp):

        project = IastProject.objects.filter(Q(id=pid),
                                             user__in=auth_users).first()

        vul = IastVulnerabilityModel.objects.filter(pk=vid).first()

        if project or vul:
            if not project:
                Project = namedtuple(
                    'Project',
                    get_model_field(IastProject,
                                    include=[
                                        'id', 'name', 'mode', 'latest_time',
                                        'vul_count', 'agent_count'
                                    ]))
                project = Project(id=0,
                                  name='NAN',
                                  mode='NAN',
                                  latest_time=time.time(),
                                  vul_count=1,
                                  agent_count=0)
            agent_ids = self.get_agents_with_project_id(project.id, auth_users)

            count_result = get_vul_count_by_agent(agent_ids, vid, user)

            levelInfo = IastVulLevel.objects.all()

            if type == 'docx':
                return self.generate_word_report(user, project, vul, count_result, levelInfo, timestamp)
                pass
            elif type == 'pdf':
                return self.generate_pdf_report(user, project, vul, count_result, levelInfo, timestamp)
            elif type == 'xlsx':
                return self.generate_xlsx_report(user, project, vul, count_result, levelInfo, timestamp)

        return None, None

    def generate_pdf_report(self, user, project, vul, count_result, levelInfo, timestamp):
        from django.template.loader import render_to_string
        import os

        levelNameArr = {}
        levelIdArr = {}
        if levelInfo:
            for level_item in levelInfo:
                levelNameArr[level_item.name_value] = level_item.id
                levelIdArr[level_item.id] = level_item.name_value

        type_summary = count_result['type_summary']

        levelCount = count_result['levelCount']

        vulDetail = count_result['vulDetail']

        levelCountArr = []
        if levelCount:
            for ind in levelCount.keys():
                levelCountArr.append(str(levelIdArr[ind]) + str(levelCount[ind]))
        levelCountStr = ",".join(levelCountArr)

        vulTypeTableBodyRows = []

        new_cells = []
        if type_summary:
            for type_item in type_summary:
                vulTypeTableBodyRow = {
                    "type_level": levelIdArr[type_item['type_level']],
                    "type_name": type_item['type_name'],
                    "type_count": str(type_item['type_count'])
                }
                vulTypeTableBodyRows.append(vulTypeTableBodyRow)

        vulTypeDetailArray = []
        if vulDetail:
            type_ind = 1
            for vul in vulDetail.keys():
                vulTypeDetail = {
                    "title": u'%s(%s)' % ("2.3." + str(type_ind) + "  " + vul, len(vulDetail[vul])),
                    "vuls": []
                }
                vulTypeDetailArray.append(vulTypeDetail)
                if vulDetail[vul]:
                    ind = 1
                    for one in vulDetail[vul]:
                        oneVul = {
                            "title": "2.3." + str(type_ind) + "." + str(ind) + "  " + one['title'],
                            "summary": _(u'Summary'),

                            "severity_level": _("Severity level"),
                            "level_id": levelIdArr[one['level_id']],
                            "first_scan_time": _("First scan time"),
                            "first_time": one['first_time'],
                            "last_scan_time": _("First scan time"),
                            "latest_time": one['first_time'],
                            "development_language": _("Development language"),
                            "language": one['language'],
                            "vulnerability_url": _("Vulnerability URL"),
                            "url": one['url'],

                            "description": _(u'Vulnerability description'),
                            "detail": "",
                        }
                        vulTypeDetail.vuls.append(
                            oneVul
                        )
                        ind = ind + 1
                        if one['detail_data']:
                            for item in one['detail_data']:
                                oneVul.detail += u'%s' % item
                type_ind = type_ind + 1

        pdf_filename = f"{MEDIA_ROOT}/reports/vul-report-{user.id}-{timestamp}.pdf"
        html_filename = f"{MEDIA_ROOT}/reports/vul-report-{user.id}-{timestamp}.html"

        rendered = render_to_string(
            './pdf.html',
            {
                "user": user,
                "project": project,
                "vul": vul,
                "count_result": count_result,
                "level_info": levelInfo,
                "time_str": time.strftime('%Y-%m-%d %H:%M:%s', time.localtime(timestamp)),
                "levelCountStr": levelCountStr,
                "vulTypeTableBodyRows": vulTypeTableBodyRows
            }
        )
        f = open(html_filename, 'w')
        f.write(rendered)
        f.close()
        os.system("cat {} | /usr/local/bin/wkhtmltopdf - {}".format(
            html_filename,
            pdf_filename
        ))
        ProjectReportExport.delete_old_files(f"{MEDIA_ROOT}/reports/")
        return "pdf", rendered

    def generate_xlsx_report(self, user, project, vul, count_result, levelInfo, timestamp):
        levelNameArr = {}
        levelIdArr = {}
        if levelInfo:
            for level_item in levelInfo:
                levelNameArr[level_item.name_value] = level_item.id
                levelIdArr[level_item.id] = level_item.name_value

        type_summary = count_result['type_summary']

        levelCount = count_result['levelCount']

        vulDetail = count_result['vulDetail']

        levelCountArr = []
        if levelCount:
            for ind in levelCount.keys():
                levelCountArr.append(str(levelIdArr[ind]) + str(levelCount[ind]))
        levelCountStr = ",".join(levelCountArr)

        vulTypeTableBodyRows = []

        new_cells = []
        if type_summary:
            for type_item in type_summary:
                vulTypeTableBodyRow = {
                    "type_level": levelIdArr[type_item['type_level']],
                    "type_name": type_item['type_name'],
                    "type_count": str(type_item['type_count'])
                }
                vulTypeTableBodyRows.append(vulTypeTableBodyRow)

        vulTypeDetailArray = []
        if vulDetail:
            type_ind = 1
            for vul in vulDetail.keys():
                vulTypeDetail = {
                    "title": u'%s' % ("2.3." + str(type_ind) + "  " + vul),
                    "vuls": []
                }
                vulTypeDetailArray.append(vulTypeDetail)
                if vulDetail[vul]:
                    ind = 1
                    for one in vulDetail[vul]:
                        oneVul = {
                            "title": "2.3." + str(type_ind) + "." + str(ind) + "  " + one['title'],
                            "summary": _(u'Summary'),

                            "severity_level": _("Severity level"),
                            "level_id": levelIdArr[one['level_id']],
                            "first_scan_time": _("First scan time"),
                            "first_time": one['first_time'],
                            "last_scan_time": _("Last scan time"),
                            "latest_time": one['last_time'],
                            "development_language": _("Development language"),
                            "language": one['language'],
                            "vulnerability_url": _("Vulnerability URL"),
                            "url": one['url'],

                            "description": _(u'Vulnerability description'),
                            "detail": "",
                        }
                        vulTypeDetail.vuls.append(
                            oneVul
                        )
                        ind = ind + 1
                        if one['detail_data']:
                            for item in one['detail_data']:
                                oneVul.detail += u'%s' % item
                type_ind = type_ind + 1
        from openpyxl import Workbook
        wb = Workbook()
        sheet1 = wb.active
        xlsx_filename = f"{MEDIA_ROOT}/reports/vul-report-{user.id}-{timestamp}.xlsx"

        sheet1['A1'] = str(_("Vulnerability type name"))
        sheet1['B1'] = str(_("Severity levels"))
        sheet1['C1'] = str(_("First scan time"))
        sheet1['D1'] = str(_("Last scan time"))
        sheet1['E1'] = str(_("Development language"))
        sheet1['F1'] = str(_("Vulnerability URL"))
        sheet1['G1'] = str(_('Vulnerability description'))
        line = 0
        for vulTypeDetail in vulTypeDetailArray:
            line += 1
            sheet1.write(line, 0, vulTypeDetail.title)
            for oneVul in vulTypeDetail.vuls:
                sheet1.append(
                    [vulTypeDetail.title, oneVul.level_id, oneVul.first_time, oneVul.latest_time, oneVul.language,
                     oneVul.url, oneVul.detail])
        wb.save(xlsx_filename)
        ProjectReportExport.delete_old_files(f"{MEDIA_ROOT}/reports/")
        return "xlsx", "11"

    @staticmethod
    def delete_old_files(path, save_seconds=10):
        for f in os.listdir(path):
            if os.stat(os.path.join(path, f)).st_mtime < time.time() - save_seconds:
                os.remove(os.path.join(path, f))
